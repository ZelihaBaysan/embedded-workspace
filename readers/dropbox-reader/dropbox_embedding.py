import re
import io
from typing import List, Sequence, Optional, Pattern
from llama_index.core import Document
from llama_index.core.schema import BaseNode
from dropbox import Dropbox
from dropbox.files import FileMetadata
import dropbox


class DropboxEmbeddingMethod:
    def __init__(
        self,
        access_token: str,
        root_path: Optional[str] = "",
    ):
        self.access_token = access_token
        self.root_path = root_path.rstrip("/")

    @staticmethod
    def customize_metadata(document: Document, data_source_id: str) -> Document:
        document.metadata.update({
            "data_source_id": data_source_id,
            "file_type": document.metadata.get("file_name", "").split(".")[-1].lower()
        })
        return document

    def _compile_patterns(self, patterns: List[str]) -> List[Pattern]:
        compiled = []
        for pattern in patterns:
            try:
                compiled.append(re.compile(pattern))
            except re.error:
                print(f"Invalid regex pattern: {pattern}")
        return compiled

    def _process_binary_file(self, content: bytes, file_ext: str) -> str:
        """Binary dosyaları işler (PDF, DOCX, XLSX)"""
        try:
            if file_ext == 'pdf':
                from pdfminer.high_level import extract_text
                return extract_text(io.BytesIO(content))
            elif file_ext == 'docx':
                from docx import Document as DocxDocument
                doc = DocxDocument(io.BytesIO(content))
                return "\n".join([para.text for para in doc.paragraphs])
            elif file_ext == 'xlsx':
                from openpyxl import load_workbook
                wb = load_workbook(io.BytesIO(content))
                text = ""
                for sheet in wb:
                    for row in sheet.iter_rows(values_only=True):
                        text += " ".join(str(cell) for cell in row if cell) + "\n"
                return text
            else:
                raise ValueError(f"Desteklenmeyen dosya uzantısı: {file_ext}")
        except Exception as e:
            print(f"Binary dosya işleme hatası: {str(e)}")
            return ""

    def apply_rules(
        self,
        documents: Sequence[Document],
        inclusion_rules: List[str],
        exclusion_rules: List[str],
    ) -> Sequence[Document]:
        compiled_exclude = self._compile_patterns(exclusion_rules)
        compiled_include = self._compile_patterns(inclusion_rules)

        filtered_docs = []
        print("\n[apply_rules] Başlangıç doküman sayısı:", len(documents))

        for doc in documents:
            file_path = doc.metadata.get("file_path", "")
            excluded = any(pattern.search(file_path) for pattern in compiled_exclude)
            included = any(pattern.search(file_path) for pattern in compiled_include) if compiled_include else True

            if excluded:
                print(f"[apply_rules] Dışlandı: {file_path}")
                continue
            if not included:
                print(f"[apply_rules] Dahil edilmedi: {file_path}")
                continue

            print(f"[apply_rules] Geçti: {file_path}")
            filtered_docs.append(doc)

        return filtered_docs

    def get_documents(self, data_source_id: str) -> List[Document]:
        dbx = Dropbox(self.access_token)
        documents = []

        try:
            if self.root_path:
                result = dbx.files_list_folder(path=self.root_path, recursive=True)
            else:
                result = dbx.files_list_folder(path="", recursive=True)
        except Exception as e:
            print(f"Error accessing Dropbox: {str(e)}")
            return documents

        # Process all files
        while True:
            for entry in result.entries:
                if not isinstance(entry, FileMetadata):
                    continue

                file_path = entry.path_display
                file_name = entry.name
                file_ext = file_name.split('.')[-1].lower() if '.' in file_name else ''

                try:
                    _, response = dbx.files_download(file_path)
                    content = response.content

                    # Text veya binary dosya işleme
                    try:
                        text_content = content.decode('utf-8-sig')  # Türkçe karakter desteği
                    except UnicodeDecodeError:
                        if file_ext in ['pdf', 'docx', 'xlsx']:
                            print(f"[get_documents] Binary dosya işleniyor: {file_path}")
                            text_content = self._process_binary_file(content, file_ext)
                            if not text_content.strip():
                                print(f"[get_documents] Boş içerik: {file_path}")
                                continue
                        elif file_ext in ['txt', 'log', 'md', 'csv']:
                            try:
                                text_content = content.decode('latin1')
                            except Exception as e:
                                print(f"[get_documents] Text dosya decode hatası: {file_path} - {str(e)}")
                                continue
                        else:
                            print(f"[get_documents] Desteklenmeyen binary dosya: {file_path}")
                            continue

                    doc = Document(
                        text=text_content,
                        metadata={
                            "file_path": file_path,
                            "file_name": file_name,
                            "file_extension": file_ext,
                            "last_modified": entry.client_modified.isoformat()
                        }
                    )
                    self.customize_metadata(doc, data_source_id)
                    documents.append(doc)

                except Exception as e:
                    print(f"[get_documents] Error processing {file_path}: {str(e)}")
                    continue

            if not result.has_more:
                break
            result = dbx.files_list_folder_continue(result.cursor)

        print(f"[get_documents] Toplam {len(documents)} dosya alındı")
        return documents

    def get_nodes(self, documents: Sequence[Document]) -> Sequence[BaseNode]:
        return []
